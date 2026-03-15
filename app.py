import os
import uuid
import json
import faiss
import numpy as np
import fitz # PyMuPDF
import google.generativeai as genai
from datetime import timedelta

from fastapi import FastAPI, UploadFile, File, BackgroundTasks, HTTPException, Depends, Form, Request, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, RedirectResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from fastapi.security import OAuth2PasswordRequestForm
from pydantic import BaseModel
from typing import List, Dict, Any

from sentence_transformers import SentenceTransformer
from sqlalchemy.orm import Session
from docx import Document as DocxDocument
from fpdf import FPDF

# Import our custom modules
import models
from database import engine, get_db
from auth import verify_password, get_password_hash, create_access_token, get_current_user, ACCESS_TOKEN_EXPIRE_MINUTES, oauth2_scheme

# Create DB Tables
models.Base.metadata.create_all(bind=engine)

# -----------------------------
# Create FastAPI App
# -----------------------------

app = FastAPI(title="AI Research Assistant Web App")

# Allow frontend connection
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Setup Templates and Static files
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")

# -----------------------------
# Configuration
# -----------------------------

DATA_DIR = "data"
INDEX_DIR = "index"
EXPORT_DIR = "exports"

EMBED_MODEL = "all-MiniLM-L6-v2"
EMBED_DIM = 384
TOP_K = 5

os.makedirs(DATA_DIR, exist_ok=True)
os.makedirs(INDEX_DIR, exist_ok=True)
os.makedirs(EXPORT_DIR, exist_ok=True)

INDEX_FILE = os.path.join(INDEX_DIR, "faiss.index")
META_FILE = os.path.join(INDEX_DIR, "metadata.json")

# -----------------------------
# Load Embedding Model & FAISS
# -----------------------------

embedder = SentenceTransformer(EMBED_MODEL)

if os.path.exists(INDEX_FILE):
    index = faiss.read_index(INDEX_FILE)
else:
    index = faiss.IndexFlatL2(EMBED_DIM)

metadata_store: Dict[str, Dict[str, Any]] = {}

if os.path.exists(META_FILE):
    with open(META_FILE, "r", encoding="utf-8") as f:
        metadata_store = json.load(f)

vector_counter = len(metadata_store)

# -----------------------------
# Request Models
# -----------------------------

class SummarizeRequest(BaseModel):
    paper_id: str
    length: str = "short"

class InsightRequest(BaseModel):
    paper_id: str
    questions: List[str]

class RawTextRequest(BaseModel):
    text: str
    filename: str = "Raw Text Input"

# -----------------------------
# Utility Functions
# -----------------------------

def extract_pdf_text(path: str):
    doc = fitz.open(path)
    pages = [page.get_text() for page in doc]
    return "\n".join(pages)

def chunk_text(text: str, chunk_size=2000, overlap=300):
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += chunk_size - overlap
    return chunks

def embed_texts(texts: List[str]):
    return embedder.encode(texts, convert_to_numpy=True)

def save_index():
    faiss.write_index(index, INDEX_FILE)
    with open(META_FILE, "w", encoding="utf-8") as f:
        json.dump(metadata_store, f)

# -----------------------------
# Gemini AI Response
# -----------------------------

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "AIzaSyAIS2V8cw-VhNhLlm5401yrEK5CRSDsGzo")
if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)

def call_llm(prompt: str):
    if not GEMINI_API_KEY:
        return "ERROR: GEMINI_API_KEY is not set."
    try:
        model = genai.GenerativeModel("gemini-2.5-flash")
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"AI Generation Error: {str(e)}"

# -----------------------------
# Retrieve & Ingest DB logic
# -----------------------------

def retrieve_chunks(query: str, paper_id: str):
    q_emb = embed_texts([query])[0]
    search_k = min(100, index.ntotal) if index.ntotal > 0 else TOP_K
    if search_k == 0:
        return []

    distances, ids = index.search(np.expand_dims(q_emb, axis=0), search_k)
    results = []

    for vid in ids[0]:
        if vid == -1: continue
        meta = metadata_store.get(str(vid))
        if meta and meta["paper_id"] == paper_id:
            results.append(meta["text"])
            if len(results) >= TOP_K:
                break
    return results

def ingest_text_data(text: str, paper_id: str):
    global vector_counter
    chunks = chunk_text(text)
    if not chunks:
        return
    embeddings = embed_texts(chunks)

    for i, emb in enumerate(embeddings):
        vector_id = str(vector_counter)
        index.add(np.expand_dims(emb, axis=0))
        metadata_store[vector_id] = {
            "paper_id": paper_id,
            "chunk_index": i,
            "text": chunks[i][:2000]
        }
        vector_counter += 1
    save_index()

def ingest_pdf(file_path: str, paper_id: str):
    text = extract_pdf_text(file_path)
    ingest_text_data(text, paper_id)

# -----------------------------
# FRONTEND ROUTES (HTML)
# -----------------------------

def get_current_user_optional(request: Request, db: Session = Depends(get_db)):
    token = request.cookies.get("access_token")
    if not token:
        return None
    try:
        token = token.replace("Bearer ", "")
        user = get_current_user(token, db)
        return user
    except:
        return None

@app.get("/", response_class=HTMLResponse)
async def home(request: Request, db: Session = Depends(get_db)):
    user = get_current_user_optional(request, db)
    return templates.TemplateResponse("index.html", {"request": request, "user": user})

@app.get("/features", response_class=HTMLResponse)
async def features(request: Request, db: Session = Depends(get_db)):
    user = get_current_user_optional(request, db)
    return templates.TemplateResponse("features.html", {"request": request, "user": user})

@app.get("/about", response_class=HTMLResponse)
async def about(request: Request, db: Session = Depends(get_db)):
    user = get_current_user_optional(request, db)
    return templates.TemplateResponse("about.html", {"request": request, "user": user})

@app.get("/contact", response_class=HTMLResponse)
async def contact(request: Request, db: Session = Depends(get_db)):
    user = get_current_user_optional(request, db)
    return templates.TemplateResponse("contact.html", {"request": request, "user": user})

@app.get("/login", response_class=HTMLResponse)
async def login_page(request: Request):
    return templates.TemplateResponse("login.html", {"request": request})

@app.get("/register", response_class=HTMLResponse)
async def register_page(request: Request):
    return templates.TemplateResponse("register.html", {"request": request})

@app.get("/dashboard", response_class=HTMLResponse)
async def dashboard(request: Request, db: Session = Depends(get_db)):
    user = get_current_user_optional(request, db)
    if not user:
        return RedirectResponse(url="/login", status_code=status.HTTP_302_FOUND)
    
    # Fetch user's documents
    user_docs = db.query(models.Document).filter(models.Document.user_id == user.id).order_by(models.Document.upload_date.desc()).all()
    
    return templates.TemplateResponse("dashboard.html", {"request": request, "user": user, "documents": user_docs})


# -----------------------------
# AUTH API ENDPOINTS
# -----------------------------

@app.post("/api/register")
def register_user(username: str = Form(...), password: str = Form(...), db: Session = Depends(get_db)):
    db_user = db.query(models.User).filter(models.User.username == username).first()
    if db_user:
        raise HTTPException(status_code=400, detail="Username already registered")
    
    hashed_password = get_password_hash(password)
    new_user = models.User(username=username, hashed_password=hashed_password)
    db.add(new_user)
    db.commit()
    db.refresh(new_user)
    
    # Auto login on register
    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = create_access_token(data={"sub": new_user.username}, expires_delta=access_token_expires)
    
    response = RedirectResponse(url="/dashboard", status_code=status.HTTP_302_FOUND)
    response.set_cookie(key="access_token", value=f"Bearer {access_token}", httponly=True)
    return response

@app.post("/api/login")
def login(username: str = Form(...), password: str = Form(...), db: Session = Depends(get_db)):
    user = db.query(models.User).filter(models.User.username == username).first()
    if not user or not verify_password(password, user.hashed_password):
        raise HTTPException(status_code=400, detail="Incorrect username or password")
    
    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = create_access_token(data={"sub": user.username}, expires_delta=access_token_expires)
    
    response = RedirectResponse(url="/dashboard", status_code=status.HTTP_302_FOUND)
    response.set_cookie(key="access_token", value=f"Bearer {access_token}", httponly=True)
    return response

@app.get("/logout")
def logout():
    response = RedirectResponse(url="/", status_code=status.HTTP_302_FOUND)
    response.delete_cookie("access_token")
    return response


# -----------------------------
# APP API ENDPOINTS
# -----------------------------

@app.post("/api/upload")
async def upload_pdf(request: Request, background_tasks: BackgroundTasks, file: UploadFile = File(...), db: Session = Depends(get_db)):
    user = get_current_user_optional(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")

    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(400, "Only PDF files allowed")

    paper_id = str(uuid.uuid4())
    save_path = os.path.join(DATA_DIR, f"{paper_id}.pdf")

    with open(save_path, "wb") as f:
        f.write(await file.read())

    # Save to DB
    new_doc = models.Document(user_id=user.id, filename=file.filename, paper_id=paper_id)
    db.add(new_doc)
    db.commit()
    
    background_tasks.add_task(ingest_pdf, save_path, paper_id)

    return {"paper_id": paper_id, "status": "processing", "filename": file.filename}


@app.post("/api/upload-text")
async def upload_text(req: RawTextRequest, request: Request, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    user = get_current_user_optional(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")

    paper_id = str(uuid.uuid4())
    
    # Save to DB
    new_doc = models.Document(user_id=user.id, filename=req.filename, paper_id=paper_id)
    db.add(new_doc)
    db.commit()
    
    background_tasks.add_task(ingest_text_data, req.text, paper_id)

    return {"paper_id": paper_id, "status": "processing", "filename": req.filename}


@app.post("/api/summarize")
def summarize(req: SummarizeRequest, request: Request, db: Session = Depends(get_db)):
    user = get_current_user_optional(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")

    doc = db.query(models.Document).filter(models.Document.paper_id == req.paper_id, models.Document.user_id == user.id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    chunks = retrieve_chunks("Summarize the research paper", req.paper_id)
    if not chunks:
        return {"summary": "Document processing or not found. Please wait a moment."}

    context = "\n".join(chunks)
    prompt = f"Summarize the research paper.\nContext:\n{context}\nCreate Abstract, Key Contributions, Methodology, and Limitations.\nLength: {req.length}"

    response = call_llm(prompt)
    
    # Save the generated summary back to DB
    doc.summary_text = response
    db.commit()

    return {"summary": response}


@app.post("/api/insights")
def extract_insights(req: InsightRequest, request: Request, db: Session = Depends(get_db)):
    user = get_current_user_optional(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")

    answers = {}
    for question in req.questions:
        if not question.strip(): continue
        chunks = retrieve_chunks(question, req.paper_id)
        prompt = f"Context:\n{chr(10).join(chunks)}\nAnswer clearly in exactly one meaningful line:\n{question}"
        answers[question] = call_llm(prompt)

    return {"answers": answers}


@app.get("/export/{paper_id}")
def export_summary(paper_id: str, format: str, request: Request, db: Session = Depends(get_db)):
    user = get_current_user_optional(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")
        
    doc = db.query(models.Document).filter(models.Document.paper_id == paper_id, models.Document.user_id == user.id).first()
    if not doc or not doc.summary_text:
        raise HTTPException(status_code=404, detail="Summary not found for export")
    
    safe_filename = "".join(c for c in doc.filename if c.isalnum() or c in (' ', '_', '-')).rstrip()
    
    if format == 'txt':
        filepath = os.path.join(EXPORT_DIR, f"{paper_id}.txt")
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(f"Summary for: {doc.filename}\n\n{doc.summary_text}")
        return FileResponse(filepath, filename=f"Summary_{safe_filename}.txt")
        
    elif format == 'docx':
        filepath = os.path.join(EXPORT_DIR, f"{paper_id}.docx")
        d = DocxDocument()
        d.add_heading(f"Summary: {doc.filename}", 0)
        d.add_paragraph(doc.summary_text)
        d.save(filepath)
        return FileResponse(filepath, filename=f"Summary_{safe_filename}.docx")
        
    elif format == 'pdf':
        filepath = os.path.join(EXPORT_DIR, f"{paper_id}.pdf")
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(200, 10, txt=f"Summary: {doc.filename[:30]}...", ln=True, align='L')
        pdf.set_font("Arial", size=11)
        # Replacing problematic unicode chars that standard FPDF fonts can't print easily
        clean_text = doc.summary_text.encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 8, txt=clean_text)
        pdf.output(filepath)
        return FileResponse(filepath, filename=f"Summary_{safe_filename}.pdf")
        
    raise HTTPException(status_code=400, detail="Invalid format requested")
